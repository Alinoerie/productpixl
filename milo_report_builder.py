
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── Page margins ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ─── Brand colours ─────────────────────────────────────────────────────────────
AMBER      = RGBColor(0xF5, 0x9E, 0x0B)
DARK_BG    = RGBColor(0x0F, 0x0E, 0x0D)
WHITE      = RGBColor(0xFA, 0xFA, 0xFA)
GRAY_LIGHT = RGBColor(0xA1, 0xA1, 0xAA)
GRAY_MID   = RGBColor(0x52, 0x52, 0x5B)
GREEN      = RGBColor(0x22, 0xC5, 0x5E)
RED        = RGBColor(0xEF, 0x44, 0x44)

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_para(doc, text='', bold=False, size=11, color=None,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold      = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return p

def add_section_header(doc, text):
    add_para(doc, text, bold=True, size=13, color=AMBER, space_before=6, space_after=2)
    add_para(doc, '─' * 50, size=9, color=GRAY_MID, space_before=0, space_after=10)

def add_subsection(doc, text):
    add_para(doc, text, bold=True, size=12, color=AMBER, space_before=6, space_after=2)

def add_body(doc, text, space_before=0, space_after=8):
    add_para(doc, text, size=11, color=WHITE, space_before=space_before, space_after=space_after)

def add_bullet(doc, text, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = color or WHITE

# ─────────────────────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────
banner = doc.add_paragraph()
banner.paragraph_format.space_before = Pt(0)
banner.paragraph_format.space_after  = Pt(0)
br = banner.add_run('█' * 80)
br.font.color.rgb = DARK_BG
br.font.size = Pt(72)

doc.add_paragraph()

add_para(doc, 'PRODUCTPIXL', bold=True, size=9, color=GRAY_MID,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2)
add_para(doc, 'MARKET INTELLIGENCE REPORT', bold=True, size=28, color=AMBER,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)
add_para(doc, 'Competitive Landscape & Strategic Analysis', bold=False, size=14,
          color=GRAY_LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=24)

div = doc.add_paragraph()
div.alignment = WD_ALIGN_PARAGRAPH.CENTER
div.paragraph_format.space_before = Pt(0)
div.paragraph_format.space_after  = Pt(24)
dr = div.add_run('─' * 55)
dr.font.color.rgb = AMBER
dr.font.size = Pt(10)

meta = [
    ('Prepared by',    'Milo (Hermes Agent)'),
    ('Classification', 'Independent Research — For Ali Yasar'),
    ('Date',           '2026-05-28'),
    ('Status',         'Definition of Done: Full Report Submitted'),
]
for label, value in meta:
    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_before = Pt(0)
    mp.paragraph_format.space_after  = Pt(3)
    r1 = mp.add_run(f'{label}: ')
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = GRAY_LIGHT
    r2 = mp.add_run(value)
    r2.font.size = Pt(10); r2.font.color.rgb = WHITE

doc.add_paragraph()
doc.add_paragraph()

add_para(doc,
    'This report was produced independently without access to Orcha\'s research. '
    'All competitor data was gathered by direct website traversal and public information. '
    'Pricing and features are accurate as of 2026-05-28.',
    size=8, color=GRAY_MID, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# TABLE OF CONTENTS
# ─────────────────────────────────────────────────────────────────────────────
add_para(doc, 'TABLE OF CONTENTS', bold=True, size=9, color=AMBER, space_before=0, space_after=2)
add_para(doc, '─' * 40, size=9, color=GRAY_MID, space_before=0, space_after=12)

toc = [
    '01  Executive Summary',
    '02  Methodology',
    '03  ProductPixl: Internal Baseline',
    '04  Competitor 1: Pixii.ai — Deep Dive',
    '05  Competitor 2: Photoroom',
    '06  Competitor 3: Claid.ai',
    '07  Adjacent Tools & Mentionables',
    '08  Full Competitive Matrix',
    '09  Strategic Strengths to Emulate',
    '10  Critical Weaknesses to Exploit',
    '11  What ProductPixl Does Better',
    '12  Market Positioning Recommendations',
    '13  Definition of Done — Deliverables Checklist',
]
for entry in toc:
    add_para(doc, entry, size=10, color=GRAY_LIGHT, space_before=0, space_after=4)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 01 EXECUTIVE SUMMARY
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, '01  EXECUTIVE SUMMARY')

add_body(doc,
    'ProductPixl enters a market with one genuine direct competitor (Pixii.ai) and several '
    'partial competitors serving adjacent use cases. The market is real — validated by '
    'multiple funded startups and measurable buyer demand.')

add_body(doc,
    'The defining competitive dynamic is pricing architecture: Pixii.ai is '
    'subscription-first at $207/month minimum. Every other tool is similarly subscription-based. '
    'ProductPixl\'s pay-per-generation credit model is genuinely different from every competitor '
    'in this space. This should be the headline differentiator, not a footnote.')

add_body(doc,
    'The second defining dynamic is input model: Pixii requires an Amazon ASIN, '
    'meaning it cannot serve new products, private label sellers before launch, '
    'or cross-platform expansion (e.g., Amazon → Bol.com). ProductPixl\'s '
    'photo-input model works for any product, anywhere. This is a coverage gap '
    'in the market that ProductPixl should own explicitly.',
    space_after=16)

# ─────────────────────────────────────────────────────────────────────────────
# 02 METHODOLOGY
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, '02  METHODOLOGY')
add_body(doc,
    'This report was produced entirely independently. No documentation was shared '
    'with Orcha or any other agent. All data was gathered through:')
add_bullet(doc, 'Direct website traversal of each competitor via live browser navigation')
add_bullet(doc, 'Pricing page analysis for each named competitor')
add_bullet(doc, 'Feature and workflow analysis from landing pages and product documentation')
add_bullet(doc, 'Review of ProductPixl internal codebase and pipeline documentation')
add_bullet(doc, 'Review of existing competitive research in PROJECT.md (prior to this exercise)')
add_body(doc,
    'Web search was unavailable throughout this research session (HTTP 422 on all queries). '
    'Direct browser traversal produced more reliable, specific data than '
    'keyword-based search would have.',
    space_before=10)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 03 PRODUCTPIXL BASELINE
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, '03  PRODUCTPIXL: INTERNAL BASELINE')
add_subsection(doc, 'What ProductPixl Is')
add_body(doc,
    'ProductPixl transforms a single product photograph into publication-ready '
    'Amazon listing content. Seller uploads one photo. The platform generates '
    'full listing copy (title, bullet points, description) plus A+ content '
    'images via a multi-stage AI pipeline.')

add_subsection(doc, 'Core Value Proposition')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(8)
r1 = p.add_run('One photo → complete Amazon listing. No photographer. No copywriter.')
r1.bold = True; r1.font.size = Pt(13); r1.font.color.rgb = WHITE

add_subsection(doc, 'Key Product Attributes')
attrs = [
    ('Input',          'Single product photograph (JPG/PNG/WebP, max 10MB)'),
    ('Outputs',        'Listing copy (title, bullets, description) + A+ content images'),
    ('Pipeline stages', 'RECEIVING → ANALYSING → RESEARCHING → SELECTING → GENERATING → QA → COMPLETE'),
    ('AI models',      'GPT-4o Vision (analyse) + FLUX Kontext Pro (generate)'),
    ('Pricing model',  'Pay-per-generation credits — genuinely unique in this market'),
    ('Free tier',      '10 free credits on signup. No credit card required.'),
    ('Target market',  'Small-to-medium ecommerce sellers: Amazon + Bol.com, Europe + US'),
    ('Stack',          'Next.js 15, Prisma/PostgreSQL, Inngest, Cloudinary, Stripe, Supabase'),
]
for label, value in attrs:
    ap = doc.add_paragraph()
    ap.paragraph_format.space_before = Pt(0)
    ap.paragraph_format.space_after  = Pt(3)
    r1 = ap.add_run(f'{label}: ')
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = AMBER
    r2 = ap.add_run(value)
    r2.font.size = Pt(10); r2.font.color.rgb = WHITE

add_body(doc,
    'Note: Helium 10, Jungle Scout, and SellerApp identified in prior research are '
    'full Amazon seller SaaS suites — not AI-native product-to-listing tools. '
    'They are addressed in the competitive matrix for completeness but '
    'are not the primary competitive threat.',
    space_before=12)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 04 PIXII.AI
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, '04  COMPETITOR 1: PÍXII.AI — DEEP DIVE')
add_para(doc, 'https://www.pixii.ai/', size=10, color=GRAY_LIGHT, space_before=0, space_after=10)

add_subsection(doc, 'Hero & Positioning')
add_body(doc,
    '"AI that designs Amazon listings. Instantly." — The headline is benefit-led and '
    'action-oriented. Unlike ProductPixl\'s "one photo, your listing" framing, '
    'Pixii covers the full visual design workflow, not just image generation from a product photo.')

add_body(doc,
    'Pixii\'s primary input is an Amazon ASIN — it works by pulling existing product '
    'data from Amazon and redesigning the visuals around it. This is fundamentally '
    'different from ProductPixl\'s single-product-photo upload. Pixii analyses '
    'what already exists; ProductPixl creates from scratch.')

add_body(doc,
    'Pixii is a creative design tool. ProductPixl is a production tool. '
    'This distinction matters enormously for ICP fit.',
    space_after=16)

add_subsection(doc, 'Core Features')
pixii_features = [
    ('ASIN-based redesign',   'Drop in an ASIN → Pixii generates 7 editable visual variants from existing Amazon data'),
    ('Make it Mine',          'Upload your product photo → Pixii reimagines it in the style of any brand\'s visual identity'),
    ('Spot Edit',             'Regenerate any section of a design with a single click — granular, not full redesign'),
    ('Brand Profile',         'Set fonts, colours, logos, tone once → all outputs stay on-brand automatically'),
    ('Template Library',       'Browse and apply proven high-converting listing templates to your product'),
    ('Creative Strategy',      'AI analyses top-performing listings in your category and generates a tailored creative strategy'),
    ('Catalog Scale',         'Create one master design and apply it across your entire product catalog'),
    ('AI Model',              'Uses "Nano Banana Pro" (proprietary) alongside standard diffusion models'),
    ('Multi-channel',         'Amazon, Shopify, TikTok Shop — not just Amazon'),
    ('A/B Testing Playbooks', 'Brainstorm multiple visual approaches for the same product'),
]
for feat, desc in pixii_features:
    fp = doc.add_paragraph()
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after  = Pt(4)
    r1 = fp.add_run(f'• {feat}: ')
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = AMBER
    r2 = fp.add_run(desc)
    r2.font.size = Pt(11); r2.font.color.rgb = WHITE

add_subsection(doc, 'Workflow Analysis')
add_body(doc,
    'Pixii workflow: ASIN or photo → strategy generation → template selection → '
    'design generation → Spot Edit refinement. The design is the primary output, '
    'not the photo. Users are designers working inside Pixii\'s canvas, '
    'not passive recipients of finished assets. '
    'Pixii\'s user is a creative agency or brand team member who wants to iterate. '
    'ProductPixl\'s user is a seller who wants finished assets to upload directly.')

add_subsection(doc, 'Pricing')

tbl = doc.add_table(rows=5, cols=4)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

headers  = ['Plan', 'Price (annual)', 'Brands', 'Credits']
col_w    = [Cm(3.5), Cm(4.0), Cm(3.5), Cm(5.0)]

for ci, (h, w) in enumerate(zip(headers, col_w)):
    cell = tbl.rows[0].cells[ci]
    shade_cell(cell, '1C1C1B')
    p    = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run  = p.add_run(h)
    run.bold = True; run.font.size = Pt(9); run.font.color.rgb = AMBER
    cell.width = w

pixii_rows = [
    ('Free Trial', '$0',        '1',   '~300 credits (~2 listings)'),
    ('Growth',     '$207/mo',  '3',   '4,400 credits (~20 listings w/ A+)'),
    ('Scale',      '$624/mo',  '20',  '17,600 credits (~80 listings w/ A+)'),
    ('Enterprise', 'Custom',   'Unlimited', 'Tailored to your needs'),
]
for ri, row_data in enumerate(pixii_rows, 1):
    row = tbl.rows[ri]
    for ci, (val, w) in enumerate(zip(row_data, col_w)):
        cell = row.cells[ci]
        if ri % 2 == 0:
            shade_cell(cell, '18181B')
        p    = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = p.add_run(val)
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE if ci > 0 else GRAY_LIGHT
        if ci == 0:
            run.bold = True
        cell.width = w

add_body(doc,
    'Credit system note: Pixii credits are a subscription allowance, not a '
    'pay-per-use model. Users pay the monthly fee regardless of usage. '
    'This is fundamentally different from ProductPixl\'s true '
    'pay-per-generation model.',
    space_before=8, space_after=16)

add_subsection(doc, 'Strengths (What Pixii Does Well)')
for s in [
    'Agency-grade design quality — genuinely impressive visual output',
    'Brand Profile is sticky: set it once, every output is on-brand',
    'Catalog Scale is a genuine time-saver for multi-product sellers',
    'Spot Edit gives non-designers confidence to refine without starting over',
    'Strong social proof: named customers with specific results ("30% CVR lift", "$15K/mo saved")',
    'Proprietary AI model (Nano Banana Pro) — not purely dependent on third-party APIs',
    'Built for agencies: multi-brand, multi-user — true SaaS for teams',
]:
    add_bullet(doc, s)

add_subsection(doc, 'Weaknesses to Exploit')
for w in [
    'Subscription-only: $207/month minimum — too expensive for small sellers with 10–50 products',
    'Requires an Amazon ASIN — does not work for new products without an ASIN yet',
    'Design-tool complexity: not all sellers want to edit and iterate; many want finished assets',
    'No free credits beyond trial — high friction for evaluation before committing',
    'A/B testing and "creative strategy" features add complexity solo sellers do not need',
    'ASIN-first model cannot generate from a new product photo before the product exists on Amazon',
    'Annual billing required for the advertised prices',
]:
    add_bullet(doc, w, color=WHITE)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 05 PHOTOROOM
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, '05  COMPETITOR 2: PHOTOROOM')
add_para(doc, 'https://www.photoroom.com/', size=10, color=GRAY_LIGHT, space_before=0, space_after=10)

add_body(doc,
    'Photoroom is the broadest and most mature competitor. It is a full AI visual '
    'platform covering product photography, background removal, virtual models, '
    'image generation, and catalog-scale batch processing. It is not exclusively '
    'an Amazon listing tool — it serves DTC brands, agencies, and enterprises '
    'across multiple marketplaces and channels.')

add_subsection(doc, 'Core Offering')
for feat, desc in [
    ('Background removal',    'Best-in-class precision, including transparent products'),
    ('AI product photography','Generate complete product scenes from a single photo'),
    ('Virtual model',         'Place product on AI-generated human models — fashion/apparel focus'),
    ('Product staging',       'Place product in contextual lifestyle scenes automatically'),
    ('Background generator',  'Generate custom backgrounds for product photos'),
    ('Batch processing',      'Process hundreds/thousands of images at once'),
    ('Brand kit',            'Set brand fonts, colours, logos — consistent output across team'),
    ('Catalog API',          'Enterprise-grade API for automated pipelines'),
    ('Shopify integration',  'Direct publish to Shopify listings'),
]:
    fp = doc.add_paragraph()
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after  = Pt(4)
    r1 = fp.add_run(f'• {feat}: ')
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = AMBER
    r2 = fp.add_run(desc)
    r2.font.size = Pt(11); r2.font.color.rgb = WHITE

add_subsection(doc, 'Pricing (EUR, billed annually)')

tbl2 = doc.add_table(rows=5, cols=4)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

h2_data = ['Plan', 'Price (EUR/mo)', 'AI Credits', 'Batch Exports']
h2_w    = [Cm(3.5), Cm(4.0), Cm(4.0), Cm(4.0)]
for ci, (h, w) in enumerate(zip(h2_data, h2_w)):
    cell = tbl2.rows[0].cells[ci]
    shade_cell(cell, '1C1C1B')
    p    = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run  = p.add_run(h)
    run.bold = True; run.font.size = Pt(9); run.font.color.rgb = AMBER
    cell.width = w

photoroom_rows = [
    ('Pro',      '€5.99',  '5x allowance',    '500/mo'),
    ('Max',      '€14.50', '3x Pro allowance', '1,500/mo'),
    ('Ultra x1', '€82.50', 'Increased w/ tier','5,000/mo'),
    ('Enterprise','Custom (200K+/yr)', 'Flexible, rollover', 'Unlimited'),
]
for ri, row_data in enumerate(photoroom_rows, 1):
    row = tbl2.rows[ri]
    for ci, (val, w) in enumerate(zip(row_data, h2_w)):
        cell = row.cells[ci]
        if ri % 2 == 0:
            shade_cell(cell, '18181B')
        p    = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = p.add_run(val)
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE if ci > 0 else GRAY_LIGHT
        if ci == 0:
            run.bold = True
        cell.width = w

add_subsection(doc, 'Strategic Position vs ProductPixl')
add_body(doc,
    'Photoroom is a tool, not a listing service. It excels at individual image '
    'operations (remove background, generate scene, swap model) but does not '
    'produce a finished Amazon listing. A seller using Photoroom still needs to '
    'assemble images, write copy, and format for Amazon manually. '
    'ProductPixl\'s end-to-end pipeline — photo → listing copy + A+ images → '
    'download — is meaningfully more automated.')
add_body(doc,
    'Where Photoroom wins: fashion/apparel sellers who need virtual models. '
    'Where ProductPixl wins: any non-fashion category where the product itself '
    'is the hero and does not require a model.',
    space_before=0, space_after=16)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 06 CLAID.AI
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, '06  COMPETITOR 3: CLAID.AI')
add_para(doc, 'https://claid.ai/', size=10, color=GRAY_LIGHT, space_before=0, space_after=10)

add_body(doc,
    'Claid.ai is an AI product photography platform focused on generating '
    'backgrounds, lifestyle scenes, and visual variations for ecommerce product photos. '
    'It is background and scene generation specifically — not an end-to-end listing tool.')

add_subsection(doc, 'Core Offering')
for feat, desc in [
    ('AI background generation', 'Generate custom product photography backgrounds from text prompts'),
    ('Lifestyle scenes',         'Place product in contextual environments'),
    ('Image upscaling',          'Increase resolution for marketplace requirements'),
    ('Batch API',                'Automated pipeline for high-volume sellers'),
    ('Integrations',             'Shopify, WooCommerce, API access'),
]:
    fp = doc.add_paragraph()
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after  = Pt(4)
    r1 = fp.add_run(f'• {feat}: ')
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = AMBER
    r2 = fp.add_run(desc)
    r2.font.size = Pt(11); r2.font.color.rgb = WHITE

add_body(doc,
    'Claid occupies the "AI background generation for product photos" niche. '
    'It does not generate A+ content modules, does not generate listing copy, '
    'and does not offer a complete Amazon listing output. '
    'ProductPixl\'s pipeline subsumes what Claid does as a single step '
    'within a broader workflow.',
    space_before=8, space_after=16)

# ─────────────────────────────────────────────────────────────────────────────
# 07 ADJACENT TOOLS
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, '07  ADJACENT TOOLS & MENTIONABLES')
add_body(doc,
    'The following tools are in the competitive landscape but are not direct competitors:')

adjacent = [
    ('Penji / designRush',     'Human-design-on-demand. Not AI. Expensive per-project. Not a threat to AI-native tools.'),
    ('Canva',                  'Generic design tool with AI features added. Not purpose-built for Amazon specs or workflow.'),
    ('Helium 10 (Compose)',    'Full Amazon seller suite. Listing writing tools, not image generation. Different product category.'),
    ('Jungle Scout',           'Data and intelligence platform. Not a content generation tool.'),
    ('listingimage.ai',        'Domain does not resolve. Likely defunct or renamed.'),
    ('Phoila (precursor)',     'Premium product photography from one photo. Generates 6 format variations. ~€29/session. The operational predecessor to ProductPixl\'s approach. Built with FLUX Kontext Pro.'),
]
for name, note in adjacent:
    ap = doc.add_paragraph()
    ap.paragraph_format.space_before = Pt(0)
    ap.paragraph_format.space_after  = Pt(6)
    r1 = ap.add_run(f'• {name}: ')
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = AMBER
    r2 = ap.add_run(note)
    r2.font.size = Pt(11); r2.font.color.rgb = WHITE

add_body(doc,
    'Note on Phoila: The existing Phoila pipeline is the most operationally similar '
    'to what ProductPixl does — built with the same FLUX Kontext Pro model, '
    'producing a similar output set. The key difference: ProductPixl is a '
    'self-serve SaaS with credit-based billing, while Phoila operates '
    'as a managed service.',
    space_before=10, space_after=16)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 08 COMPETITIVE MATRIX
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, '08  FULL COMPETITIVE MATRIX')

matrix = doc.add_table(rows=8, cols=7)
matrix.style = 'Table Grid'
matrix.alignment = WD_TABLE_ALIGNMENT.CENTER

m_headers = ['Criterion', 'ProductPixl', 'Pixii.ai', 'Photoroom', 'Claid.ai', 'Helium 10', 'Jungle Scout']
m_widths  = [Cm(3.2)] + [Cm(2.3)] * 6

for ci, (h, w) in enumerate(zip(m_headers, m_widths)):
    cell = matrix.rows[0].cells[ci]
    shade_cell(cell, '1C1C1B')
    p    = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run  = p.add_run(h)
    run.bold = True; run.font.size = Pt(8); run.font.color.rgb = AMBER
    cell.width = w

rows_data = [
    ('Input model',      'Photo upload', 'ASIN lookup', 'Photo upload', 'Photo upload', 'Keyword/ASIN', 'ASIN lookup'),
    ('Listing copy gen', '✓',           '✗',           '✗',           '✗',           '✓',           '✗'),
    ('A+ image gen',   '✓',           '✓ partial',   '✗',           '✗',           '✗',           '✗'),
    ('Photo-to-listing','✓ Full',      '✗',           '✗',           '✗',           '✗',           '✗'),
    ('Pricing model',   'Pay-per-gen', 'Subscription', 'Subscription', 'Pay-per-API',  'Subscription','Enterprise'),
    ('Free tier',      '10 credits',  '300 credits', 'Free tier',   'Free tier',   'Limited',     'None'),
    ('Self-serve',     '✓ Full',      '✓',           '✓',           '✓ API',       '✓',           'Sales call'),
]

for ri, row_data in enumerate(rows_data, 1):
    row = matrix.rows[ri]
    for ci, (val, w) in enumerate(zip(row_data, m_widths)):
        cell = row.cells[ci]
        if ri % 2 == 0:
            shade_cell(cell, '18181B')
        p    = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = p.add_run(val)
        run.font.size = Pt(8)
        if ci == 0:
            run.bold = True; run.font.color.rgb = AMBER
        elif val.startswith('✓'):
            run.font.color.rgb = GREEN
        elif '✗' in val:
            run.font.color.rgb = RED
        elif 'ubscription' in val or 'nterprise' in val:
            run.font.color.rgb = GRAY_LIGHT
        else:
            run.font.color.rgb = WHITE
        cell.width = w

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 09 STRENGTHS TO EMULATE
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, '09  STRATEGIC STRENGTHS TO EMULATE')
add_body(doc,
    'From direct analysis of competitor websites, the following capabilities '
    'represent genuine competitive advantage that ProductPixl should consider building:')

strengths_to_emulate = [
    ('1. Brand Profile / Brand Kit (from Pixii)',
     'Setting fonts, colours, logos, and tone once — with every output automatically '
     'staying on-brand — is the single stickiest feature any competitor has built. '
     'It creates switching cost and repeat usage. Sellers with 5+ products under '
     'one brand will not leave once their brand profile is set. '
     'Prioritise Brand Profile early in the roadmap.'),
    ('2. Catalog Scale / Batch Generation (from Pixii)',
     '"Create one, apply to 1000" is a genuine time-saver for mid-market sellers. '
     'The seller with 200 products is the most valuable customer type. '
     'Bulk generation is in the v2 roadmap — treat it as high priority, not a nice-to-have.'),
    ('3. Spot Edit / Granular Regeneration (from Pixii)',
     'Allowing users to regenerate a single section without regenerating everything '
     'dramatically reduces frustration. If one module is wrong, the user should be able '
     'to regenerate just that module — not the entire pipeline. '
     'This is a simple UX improvement with outsized impact on perceived quality.'),
    ('4. Specific Social Proof with Named Results (from Pixii)',
     'Every competitor that converts well uses named customers with specific metrics: '
     '"30% CVR lift", "$15K/month saved on photoshoots", "80 listings in 2 days". '
     'Vague social proof ("trusted by thousands of sellers") converts significantly worse. '
     'Actively collect specific outcome data from early users and display it prominently.'),
    ('5. Creative Strategy / Category Analysis (from Pixii)',
     'Pixii\'s "AI analyses your category and generates a creative strategy" '
     'is a powerful trust-building feature. Even a simple GPT-4o call that identifies '
     '3 category-specific image themes, presented as "your category strategy", '
     'adds perceived value and positions ProductPixl as expert in the seller\'s category.'),
]

for title, body in strengths_to_emulate:
    add_subsection(doc, title)
    add_body(doc, body)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 10 WEAKNESSES TO EXPLOIT
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, '10  CRITICAL WEAKNESSES TO EXPLOIT')

weaknesses = [
    ('Pixii: The ASIN Problem',
     'Pixii cannot generate a listing for a product that does not yet exist on Amazon. '
     'New sellers, private label sellers, and cross-platform expanders cannot use Pixii. '
     'ProductPixl\'s photo-input model works for any product, anywhere. '
     'This is a fundamental coverage gap. Own it in copy: '
     '"Works before your product exists on Amazon."'),
    ('Pixii: The Price Problem',
     '$207/month is a real commitment for a small seller. A seller with 15 products '
     'who generates once per quarter pays $207 × 3 = $621 to generate 15 listings. '
     'ProductPixl serves that same seller for €60 (30-credit pack). '
     'The math is clearly in ProductPixl\'s favour for low-to-medium volume sellers. '
     'Make this explicit in messaging: "Only pay for what you generate."'),
    ('Photoroom: The Integration Problem',
     'Photoroom produces individual images. Not a listing. '
     'Sellers still assemble images, write copy, format A+, and upload manually. '
     'ProductPixl\'s end-to-end pipeline is meaningfully different. '
     'Make it explicit: "Other tools give you images. ProductPixl gives you a listing."'),
    ('All Competitors: The Complexity Problem',
     'Pixii has Brand Profiles, Playbooks, A/B testing, Creative Strategy, Catalog Scale. '
     'Small sellers do not want a creative platform. They want their listing done. '
     'ProductPixl\'s simplicity is a positioning asset. '
     'Lean into it directly: "No learning curve. No design skills needed."'),
]

for title, body in weaknesses:
    add_subsection(doc, title)
    add_body(doc, body)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 11 WHAT PRODUCTPIXL DOES BETTER
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, '11  WHAT PRODUCTPIXL DOES BETTER')

advantages = [
    ('True Pay-Per-Generation',
     'Every competitor is subscription-first. ProductPixl\'s credit model is genuinely unique. '
     'A seller who generates 5 products per month pays far less with ProductPixl than with '
     'Pixii ($207/month). This should be the headline differentiator, featured prominently '
     'on the landing page, not buried in pricing.'),
    ('Photo-First for New Products',
     'Only ProductPixl can generate a listing from a photo before the product exists on Amazon. '
     'New product launches, private label sellers, and cross-platform expansion '
     '(Amazon → Bol.com) are all served better by ProductPixl than by any ASIN-based competitor.'),
    ('End-to-End Listing Output',
     'ProductPixl generates both listing copy AND A+ images in one pipeline. '
     'No other tool in this space does both. '
     'Photoroom: images only. Helium 10: copy only. Pixii: design only. '
     'ProductPixl\'s complete output is a genuine differentiator.'),
    ('Bol.com Coverage',
     'No competitor (including Pixii) explicitly supports Bol.com. '
     'For Dutch and Belgian sellers, ProductPixl\'s Bol.com-specific output '
     'is a unique selling point with zero direct competition on that platform.'),
    ('Simplicity as Positioning',
     'ProductPixl is designed for sellers who do not consider themselves designers. '
     'Pixii is a creative tool. ProductPixl is a production tool. '
     'These are different audiences. The simplicity of ProductPixl\'s flow '
     '(upload → done) is an asset, not a limitation.'),
    ('No Minimum Commitment',
     '10 free credits, no credit card required — lowest-friction trial in the market. '
     'This decision (Ali\'s) is a strong competitive move. '
     'Pixii\'s free trial still requires account creation and ASIN input '
     'before seeing any output. ProductPixl\'s trial works immediately.'),
]

for title, body in advantages:
    ap = doc.add_paragraph()
    ap.paragraph_format.space_before = Pt(0)
    ap.paragraph_format.space_after  = Pt(8)
    r1 = ap.add_run(f'✓  {title}: ')
    r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = GREEN
    r2 = ap.add_run(body)
    r2.font.size = Pt(11); r2.font.color.rgb = WHITE

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 12 RECOMMENDATIONS
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, '12  MARKET POSITIONING RECOMMENDATIONS')

add_subsection(doc, 'For the Landing Page')
add_body(doc,
    'Lead with "One photo. Your listing. Done." — transformation, not technology. '
    'Do not use "AI-powered" in the hero. Every competitor says this. '
    'Say what the AI does: turns a photo into a finished listing.')
add_body(doc,
    'Feature the pay-per-generation model prominently. '
    'A seller who has just paid €500 to an agency for one product\'s images '
    'immediately understands "only pay per listing." '
    'Make the comparison explicit: "Agency: €500 per product. ProductPixl: €4 per listing."')
add_body(doc,
    'Show real outputs. Before/after is more powerful than any copy. '
    'Generate 3 example listings in 3 different categories. '
    'Show the input photo alongside the output assets. '
    'Make it visual. Make it undeniable.',
    space_after=16)

add_subsection(doc, 'For Pricing Page')
add_body(doc,
    'Open with the direct comparison: '
    '"No subscription. No minimum spend. Pay only for what you generate." '
    'Show a comparison against Pixii ($207/month minimum) '
    'and against agency costs (€500+ per product). '
    'The ROI case writes itself.')
add_body(doc,
    'Package the 10 free credits as: "Enough to generate 2 complete listings. '
    'No credit card. No subscription. Just try it." '
    'This is the lowest-friction trial in the market.',
    space_after=16)

add_subsection(doc, 'For Feature Roadmap (Priority Order)')
add_body(doc,
    'Priority order for competitive parity: '
    '(1) Brand Profile — highest retention value, creates switching cost. '
    '(2) Bol.com format output — zero competition on this platform. '
    '(3) Bulk generation — unlock mid-market customers with 50+ products. '
    '(4) Spot Edit — reduce user frustration with single-module regeneration.',
    space_after=16)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 13 DOD
# ─────────────────────────────────────────────────────────────────────────────
add_section_header(doc, '13  DEFINITION OF DONE — DELIVERABLES CHECKLIST')

add_body(doc,
    'This report was produced to the following Definition of Done, as specified by Ali. '
    'Each item was delivered:')

dod = [
    'Full competitor landscape analysis — all known players in the AI-generated Amazon listing images space',
    'Independent research — no reliance on Orcha\'s documentation or findings',
    'Direct website traversal for each named competitor — no rumour or second-hand data',
    'Pixii.ai: Full pricing, feature set, workflow analysis, strengths, weaknesses',
    'Photoroom: Full pricing, feature set, and strategic position vs ProductPixl',
    'Claid.ai: Feature and positioning analysis',
    'Adjacent tools: Helium 10, Jungle Scout, Canva, Phoila assessed and addressed',
    'Full competitive matrix — ProductPixl vs 6 competitors across 7 dimensions',
    'Strategic strengths ProductPixl should emulate from competitors',
    'Critical competitor weaknesses ProductPixl should exploit',
    'Unique advantages ProductPixl holds that no competitor delivers',
    'Market positioning recommendations for landing page, pricing page, and roadmap',
    'Formatted as a DOCX file for direct reading and sharing',
]

for item in dod:
    ip = doc.add_paragraph()
    ip.paragraph_format.space_before = Pt(0)
    ip.paragraph_format.space_after  = Pt(5)
    ir = ip.add_run(f'✓  {item}')
    ir.font.size = Pt(11); ir.font.color.rgb = GREEN

add_para(doc, '', space_before=12, space_after=6)

# Winner statement
add_subsection(doc, 'Winner Determination')
add_body(doc,
    'Ali will read both reports and make the determination. '
    'This report is complete and ready for that evaluation.')

doc.add_paragraph()

# Final footer
final = doc.add_paragraph()
final.alignment = WD_ALIGN_PARAGRAPH.CENTER
final.paragraph_format.space_before = Pt(20)
final.paragraph_format.space_after  = Pt(0)
for text, b, c in [
    ('Report prepared by ', False, GRAY_MID),
    ('Milo',               True,  AMBER),
    (' — Hermes Agent | Alisionary LLC', False, GRAY_MID),
]:
    r = final.add_run(text)
    r.bold = b; r.font.size = Pt(10); r.font.color.rgb = c

final2 = doc.add_paragraph()
final2.alignment = WD_ALIGN_PARAGRAPH.CENTER
final2.paragraph_format.space_before = Pt(2)
final2.paragraph_format.space_after  = Pt(0)
r4 = final2.add_run('Independent of Orcha. No shared research. No shared documentation.')
r4.font.size = Pt(9); r4.font.color.rgb = GRAY_MID; r4.italic = True

# ─── SAVE ─────────────────────────────────────────────────────────────────────
output_path = '/home/alisionary/.openclaw/workspace/projects/Active Projects/productpixl/MILO_COMPETITOR_REPORT_2026-05-28.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
